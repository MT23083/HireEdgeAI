# scripts/make_templates.py
"""
Creates 4 DOCX templates with Jinja placeholders for docxtpl:
 - resume_template.docx
 - sop_template.docx
 - cover_letter_template.docx
 - visa_cover_letter_template.docx (embassy format kept strict)

Run once:
  python scripts/make_templates.py
"""

from __future__ import annotations
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches, RGBColor

ROOT = Path(__file__).resolve().parents[1]
TEMPLATES_DIR = ROOT / "templates"
TEMPLATES_DIR.mkdir(parents=True, exist_ok=True)

def add_heading(doc: Document, text: str, size: int = 14):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    p.space_before = Pt(6); p.space_after = Pt(2)
    return p

def add_small_gray(doc: Document, text: str, size: int = 9, align_center=False):
    p = doc.add_paragraph()
    if align_center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor(120, 120, 120)
    return p

def add_bullet(doc: Document, text: str, size: int = 11):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    r.font.size = Pt(size)
    return p

def watermark_header(section, text="{{ watermark }}"):
    header = section.header
    p = header.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(180, 180, 180)

# -------- Resume --------
def make_resume_template(path: Path):
    doc = Document()
    s = doc.sections[0]
    s.top_margin = Inches(0.5); s.bottom_margin = Inches(0.5)
    s.left_margin = Inches(0.5); s.right_margin = Inches(0.5)
    watermark_header(s)

    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nrun = name.add_run("{{ full_name }}")
    nrun.bold = True; nrun.font.size = Pt(16)

    line = doc.add_paragraph(); line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for label, key in [("Email: ", "email"), ("Phone: ", "phone"),
                       ("Location: ", "location"), ("LinkedIn: ", "linkedin"), ("GitHub: ", "github")]:
        rr = line.add_run(label); rr.bold = True
        line.add_run(f"{{{{ {key} }}}}  ")

    add_small_gray(doc, "Target: {{ target_role }}", align_center=True)

    add_heading(doc, "SUMMARY"); doc.add_paragraph("{{ summary }}")
    add_heading(doc, "SKILLS"); doc.add_paragraph("{{ skills }}")

    add_heading(doc, "EXPERIENCE")
    add_small_gray(doc, "{% for e in experience_items %}")
    p = doc.add_paragraph("{{ e.title }} — {{ e.company }} ({{ e.dates }})")
    p.runs[0].bold = True
    add_small_gray(doc, "{% for b in e.bullets %}"); add_bullet(doc, "{{ b }}"); add_small_gray(doc, "{% endfor %}")
    add_small_gray(doc, "{% endfor %}")

    add_heading(doc, "PROJECTS")
    add_small_gray(doc, "{% for p in project_items %}")
    p = doc.add_paragraph("{{ p.name }} — {{ p.stack }}"); p.runs[0].bold = True
    add_small_gray(doc, "{% for b in p.bullets %}"); add_bullet(doc, "{{ b }}"); add_small_gray(doc, "{% endfor %}")
    add_small_gray(doc, "{% endfor %}")

    add_heading(doc, "EDUCATION")
    add_small_gray(doc, "{% for ed in education_items %}")
    doc.add_paragraph("{{ ed.degree }} — {{ ed.institute }} — {{ ed.score }} — {{ ed.year }}")
    add_small_gray(doc, "{% endfor %}")

    doc.save(path)

# -------- SOP --------
def make_sop_template(path: Path):
    doc = Document()
    s = doc.sections[0]
    s.top_margin = Inches(0.7); s.bottom_margin = Inches(0.7)
    s.left_margin = Inches(0.7); s.right_margin = Inches(0.7)
    watermark_header(s)

    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = t.add_run("Statement of Purpose"); rr.bold = True; rr.font.size = Pt(14)

    m = doc.add_paragraph(); m.alignment = WD_ALIGN_PARAGRAPH.CENTER
    m.add_run("{{ full_name }}").bold = True; m.add_run("  ·  "); m.add_run("{{ email }}")

    m2 = doc.add_paragraph(); m2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    m2.add_run("Program: ").bold = True; m2.add_run("{{ program }}")
    m2.add_run("   |   University: ").bold = True; m2.add_run("{{ university }}")

    doc.add_paragraph()
    doc.add_paragraph("{{ body }}")
    doc.save(path)

# -------- Job Cover Letter --------
def make_cover_letter_template(path: Path):
    doc = Document()
    s = doc.sections[0]
    s.top_margin = Inches(0.7); s.bottom_margin = Inches(0.7)
    s.left_margin = Inches(0.8); s.right_margin = Inches(0.8)
    watermark_header(s)

    head = doc.add_paragraph()
    head.add_run("{{ full_name }}").bold = True; head.add_run(" | "); head.add_run("{{ email }}")

    doc.add_paragraph()
    doc.add_paragraph("Dear Hiring Manager at {{ company }},")
    add_small_gray(doc, "Application for {{ role }}", size=10)
    doc.add_paragraph("{{ body }}")
    doc.add_paragraph("Sincerely,"); doc.add_paragraph("{{ full_name }}")
    doc.save(path)

# -------- Visa Cover Letter (Embassy format) --------
def make_visa_cover_letter_template(path: Path):
    doc = Document()
    s = doc.sections[0]
    s.top_margin = Inches(0.7); s.bottom_margin = Inches(0.7)
    s.left_margin = Inches(0.8); s.right_margin = Inches(0.8)
    watermark_header(s)

    # 1) Applicant block
    doc.add_paragraph("{{ applicant_block }}")
    doc.add_paragraph("{{ date }}")

    # 2) Embassy block
    p = doc.add_paragraph("{{ embassy_block }}"); p.paragraph_format.space_after = Pt(12)

    # 3) Subject
    subj = doc.add_paragraph(); r = subj.add_run("Subject: "); r.bold = True; subj.add_run("{{ subject }}")

    doc.add_paragraph()  # spacer

    # 4) Salutation
    doc.add_paragraph("Respected Visa Officer,")

    # 5) Body
    body = doc.add_paragraph("{{ body }}"); body.paragraph_format.space_after = Pt(12)

    # 6) Closing
    doc.add_paragraph("Thank you for your time and consideration.")
    doc.add_paragraph()
    doc.add_paragraph("Yours sincerely,")
    doc.add_paragraph()
    doc.add_paragraph("{{ sign_name }}")

    doc.save(path)

if __name__ == "__main__":
    make_resume_template(TEMPLATES_DIR / "resume_template.docx")
    make_sop_template(TEMPLATES_DIR / "sop_template.docx")
    make_cover_letter_template(TEMPLATES_DIR / "cover_letter_template.docx")
    make_visa_cover_letter_template(TEMPLATES_DIR / "visa_cover_letter_template.docx")
    print(f"Templates created in: {TEMPLATES_DIR}")
